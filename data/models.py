from django.db.models.signals import post_save
from django.dispatch import receiver
from django.db import models
from random import shuffle

# Create your models here.
def get_random_texts(num=24):
	texts=[]
	while len(texts) < num:
		texts += list(Text.objects.all()[:num])#TODO shuffle
	shuffle(texts)
	return texts[:num]

class Text(models.Model):
	text = models.CharField(max_length=255)

	def __str__(self):
		return self.text

class Game(models.Model):
	label = models.CharField(max_length=255, default='GAAAME')
	timestamp_add = models.DateTimeField(auto_now_add=True)
	timestamp_edit = models.DateTimeField(auto_now=True)

	def __str__(self):
		return "G {} {}".format(self.label, self.timestamp_add)

	def emit(self, fn):
		import docx
		document = docx.Document()
		document.add_heading(self.label, 0)
		table = document.add_table(rows=5, cols=5)
		for p in self.placement_set.all():
			assert 0 <= p.position < 25
			assert p.position != 12
			x,y = int(p.position/5), p.position%5
			table.rows[x].cells[y].text = p.text.text
		table.rows[2].cells[2].text = 'BINGO'
		document.save(fn)

	def init(self):
		for i,t in enumerate(get_random_texts()):
			if i >=12:
				i += 1
			Placement(game=self, text=t, position=i).save()

class Placement(models.Model):
	text = models.ForeignKey(Text, on_delete=models.CASCADE)
	game = models.ForeignKey(Game, on_delete=models.CASCADE)
	position = models.SmallIntegerField()

	class Meta:
		unique_together = (('text', 'game', 'position'),)

@receiver(post_save, sender=Game)
def game_post_save(sender, instance, created, **kwargs):
	if created:
		instance.init()
